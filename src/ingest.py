import os
from pathlib import Path
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document as LangchainDocument
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_pinecone import PineconeVectorStore

# 1. Load environment variables
load_dotenv()

# We create a small helper class to force the dimensions to 512
class SlicedGeminiEmbeddings(GoogleGenerativeAIEmbeddings):
    def embed_documents(self, texts):
        vectors = super().embed_documents(texts)
        return [v[:512] for v in vectors]

    def embed_query(self, text):
        vector = super().embed_query(text)
        return vector[:512]


def load_document(file_path: str):
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()

    if ext == ".txt":
        loader = TextLoader(file_path, encoding="utf-8")
        return loader.load()

    if ext == ".docx":
        document = DocxDocument(file_path)
        text = "\n".join([p.text for p in document.paragraphs if p.text])
        return [LangchainDocument(page_content=text, metadata={"source": file_path})]

    raise ValueError("Unsupported file type. Upload .pdf, .txt, or .docx files.")


def split_documents(documents):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return splitter.split_documents(documents)


def ingest_document(file_path: str, index_name: str | None = None):
    index_name = index_name or os.getenv("PINECONE_INDEX_NAME")
    if not index_name:
        raise ValueError("PINECONE_INDEX_NAME must be set.")

    documents = load_document(file_path)
    chunks = split_documents(documents)
    print(f"Prepared {len(chunks)} chunks for {file_path}.")

    embeddings = SlicedGeminiEmbeddings(
        model="gemini-embedding-2",
        task_type="retrieval_document"
    )

    PineconeVectorStore.from_documents(
        documents=chunks,
        embedding=embeddings,
        index_name=index_name
    )
    print("Success! Ingestion complete.")


def ingest_docs():
    ingest_document("./data/ebook.pdf")


if __name__ == "__main__":
    ingest_docs()
