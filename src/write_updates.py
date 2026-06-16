from pathlib import Path

main_content = '''import os
import tempfile
from pathlib import Path
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from dotenv import load_dotenv
from src.graph import app as graph_app
from src.ingest import ingest_document

load_dotenv()

app = FastAPI(title="Agentic AI RAG API")

class Query(BaseModel):
    question: str

@app.get("/")
def read_root():
    return {"status": "active", "info": "Agentic AI RAG System"}

@app.post("/ask")
async def ask_question(query: Query):
    # This calls your LangGraph
    inputs = {"question": query.question}
    result = await graph_app.ainvoke(inputs)
    
    return {
        "answer": result["answer"],
        "confidence": result["score"],
        "context": result["context"]
    }

@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    suffix = Path(file.filename).suffix.lower()
    if suffix not in [".pdf", ".txt", ".docx"]:
        return JSONResponse(
            status_code=400,
            content={"error": "Unsupported file type. Upload .pdf, .txt, or .docx."}
        )

    tmp_path = None
    try:
        contents = await file.read()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp_path = tmp.name
            tmp.write(contents)

        ingest_document(tmp_path)
        return {"status": "success", "filename": file.filename}
    except Exception as exc:
        return JSONResponse(status_code=500, content={"error": str(exc)})
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
'''

ui_content = '''import os
import requests
import gradio as gr

API_URL = os.getenv("API_URL", "http://127.0.0.1:8000/ask")
UPLOAD_URL = os.getenv("UPLOAD_URL", "http://127.0.0.1:8000/upload")


def upload_document(file):
    if not file:
        return "No file selected."

    path = file if isinstance(file, str) else file.name
    try:
        with open(path, "rb") as f:
            files = {"file": (os.path.basename(path), f)}
            response = requests.post(UPLOAD_URL, files=files)
            response.raise_for_status()
            data = response.json()
            return f"Uploaded: {data.get('filename')}"
    except Exception as e:
        return f"Upload failed: {e}"


def ask_question(message, history):
    if not message:
        return history or []

    try:
        response = requests.post(API_URL, json={"question": message})
        response.raise_for_status()
        data = response.json()
        answer = data.get("answer", "No answer received.")
    except Exception as e:
        answer = f"Error: Could not connect to the AI server. ({e})"

    history = history or []
    history.append(("You", message))
    history.append(("Bot", answer))
    return history


demo = gr.Blocks()
with demo:
    gr.Markdown("# Agentic AI Document QA")
    gr.Markdown("Upload a PDF, TXT, or DOCX file and ask questions about its content.")

    with gr.Row():
        file_input = gr.File(label="Upload document", file_types=[".pdf", ".txt", ".docx"])
        upload_button = gr.Button("Upload Document")
    upload_status = gr.Textbox(label="Upload Status", interactive=False)

    chatbot = gr.Chatbot(label="Chat History")
    question = gr.Textbox(label="Ask a question", placeholder="Ask anything about the uploaded document...")
    ask_button = gr.Button("Send question")

    upload_button.click(upload_document, inputs=file_input, outputs=upload_status)
    ask_button.click(ask_question, inputs=[question, chatbot], outputs=chatbot)
    question.submit(ask_question, inputs=[question, chatbot], outputs=chatbot)

if __name__ == "__main__":
    demo.launch(server_name="0.0.0.0", server_port=7860, share=False)
'''

ingest_content = '''import os
from pathlib import Path
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader
from langchain.document_loaders import TextLoader
from langchain.schema import Document as LangchainDocument
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_pinecone import PineconeVectorStore

# 1. Load environment variables
load_dotenv()

# We create a small helper class to force the dimensions to 512
class SlicedGeminiEmbeddings(GoogleGenerativeAIEmbeddings):
    def embed_documents(self, texts):
        vectors = super().embed_documents(texts)
        return [v[:512] for v in vectors]

    def embed_query(self, text):
        vector = super().embed_query(text)
        return vector[:512]


def load_document(file_path: str):
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()

    if ext == ".txt":
        loader = TextLoader(file_path, encoding="utf-8")
        return loader.load()

    if ext == ".docx":
        document = DocxDocument(file_path)
        text = "\n".join([p.text for p in document.paragraphs if p.text])
        return [LangchainDocument(page_content=text, metadata={"source": file_path})]

    raise ValueError("Unsupported file type. Upload .pdf, .txt, or .docx files.")


def split_documents(documents):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return splitter.split_documents(documents)


def ingest_document(file_path: str, index_name: str | None = None):
    index_name = index_name or os.getenv("PINECONE_INDEX_NAME")
    if not index_name:
        raise ValueError("PINECONE_INDEX_NAME must be set.")

    documents = load_document(file_path)
    chunks = split_documents(documents)
    print(f"Prepared {len(chunks)} chunks for {file_path}.")

    embeddings = SlicedGeminiEmbeddings(
        model="gemini-embedding-2",
        task_type="retrieval_document"
    )

    PineconeVectorStore.from_documents(
        documents=chunks,
        embedding=embeddings,
        index_name=index_name
    )
    print("Success! Ingestion complete.")


def ingest_docs():
    ingest_document("./data/ebook.pdf")


if __name__ == "__main__":
    ingest_docs()
'''

Path('src/main.py').write_text(main_content, encoding='utf-8')
Path('src/ui.py').write_text(ui_content, encoding='utf-8')
Path('src/ingest.py').write_text(ingest_content, encoding='utf-8')
print('wrote')
